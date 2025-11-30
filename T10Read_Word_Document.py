# Name:     Read_Word_Document.py
# Function: To show simple error trapping
# Author:   Oscar Correia
# Date:     26.03.2025

import docx

doc = docx.Document('This is a test.docx')
print(len(doc.paragraphs))

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText('This is a test.docx'))